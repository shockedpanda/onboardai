import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from watch_and_process import clean_address

def safe_str(value, default="N/A"):
    """Return string representation or default if None."""
    if value is None:
        return default
    return str(value)

def generate_kyc_report(json_path: str, output_path: str):
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    title = doc.add_heading('KYC Compliance Review Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    applicant = (data.get("individual_applicants") or [{}])[0] if data.get("individual_applicants") else {}
    
    doc.add_heading('Applicant Details', level=1)
    table = doc.add_table(rows=8, cols=2, style='Light Shading Accent 1')
    for row in table.rows:
        row.cells[0].width = Inches(1.5)

    # Clean address
    raw_address = applicant.get("residential_address", "")
    if isinstance(raw_address, dict):
        raw_address = ", ".join(str(v) for v in raw_address.values() if v)
    address_clean = clean_address(raw_address) if raw_address else ""

    pep_decl = applicant.get("pep_declaration")
    if pep_decl is True:
        pep_str = "Yes"
    elif pep_decl is False:
        pep_str = "No"
    else:
        pep_str = "Not stated"

    cells = [
        ("Full Name", safe_str(applicant.get("full_name"))),
        ("Date of Birth", safe_str(applicant.get("date_of_birth"))),
        ("Nationality", safe_str(applicant.get("nationality"))),
        ("Passport/ID No.", safe_str(applicant.get("passport_or_id_number"))),
        ("Residential Address", address_clean or "N/A"),
        ("PEP Declaration", pep_str),
        ("Source of Funds", safe_str(applicant.get("source_of_funds"))),
        ("Source of Wealth", safe_str(applicant.get("source_of_wealth")))
    ]
    for i, (label, value) in enumerate(cells):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    # Risk assessment
    doc.add_heading('Risk Assessment', level=1)
    edd = data.get("final_edd", False)
    p = doc.add_paragraph()
    p.add_run("Enhanced Due Diligence (EDD) required: ").bold = True
    p.add_run("YES" if edd else "NO")

    # Sanctions & PEP Screening
    doc.add_heading('Sanctions & PEP Screening', level=1)
    matches = data.get("sanctions_screening") or []
    if matches:
        sanctions_matches = [m for m in matches if m.get("source_type") == "sanction"]
        pep_matches = [m for m in matches if m.get("source_type") == "pep"]
        if sanctions_matches:
            doc.add_paragraph(f"⚑ Sanctions hits: {len(sanctions_matches)}", style='Intense Quote')
            tbl = doc.add_table(rows=1, cols=2, style='Medium Shading 1 Accent 1')
            tbl.rows[0].cells[0].text = 'Matched Name'
            tbl.rows[0].cells[1].text = 'Score'
            for m in sanctions_matches:
                row = tbl.add_row().cells
                row[0].text = safe_str(m.get("matched_name"))
                row[1].text = safe_str(m.get("score"))
        if pep_matches:
            doc.add_paragraph(f"⚑ PEP hits: {len(pep_matches)}", style='Intense Quote')
            tbl = doc.add_table(rows=1, cols=2, style='Medium Shading 1 Accent 1')
            tbl.rows[0].cells[0].text = 'Matched Name'
            tbl.rows[0].cells[1].text = 'Score'
            for m in pep_matches:
                row = tbl.add_row().cells
                row[0].text = safe_str(m.get("matched_name"))
                row[1].text = safe_str(m.get("score"))
    else:
        doc.add_paragraph("No sanctions or PEP matches were found.")

    # Document Completeness
    completeness = data.get("completeness") or {}
    doc.add_heading('Document Completeness', level=1)
    present = completeness.get("documents_present") or []
    missing = completeness.get("documents_missing") or []
    doc.add_paragraph(f"Documents present: {', '.join(present) if present else 'None'}")
    doc.add_paragraph(f"Documents missing: {', '.join(missing) if missing else 'None'}")

    # Reviewer Notes
    doc.add_heading('Reviewer Notes', level=1)
    doc.add_paragraph(data.get("notes") or "No additional notes.")

    doc.save(output_path)
    print(f"Report saved to {output_path}")